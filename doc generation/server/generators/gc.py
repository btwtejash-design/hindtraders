import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from typing import Dict, Any

def generate_gc_docx(data: Dict[str, Any], output_path: str) -> str:
    """
    Generates Warranty / Guarantee Certificate Word (.docx) file matching GC.png format.
    Uses dynamic Consignee from PO for the "To" block.
    """
    doc = docx.Document()

    # Set 0.75 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # 1. Company Header
    candidate_paths = [
        os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "sample", "picture.png")),
        os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "sample", "picture.png")),
        os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "client", "src", "assets", "picture.png")),
        os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "client", "public", "picture.png")),
    ]
    picture_path = None
    for p in candidate_paths:
        if os.path.exists(p):
            picture_path = p
            break

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if picture_path and os.path.exists(picture_path):
        run_title = p_title.add_run()
        run_title.add_picture(picture_path, height=Inches(0.55))
    else:
        run_title = p_title.add_run("Hind Traders")
        run_title.font.name = "Georgia"
        run_title.font.size = Pt(28)
        run_title.font.italic = True
        run_title.font.bold = True

    p_addr = doc.add_paragraph()
    p_addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_addr = p_addr.add_run("Address: - Chhoti Keshopur Nakki Nagar Jamalpur, Bihar - 811214")
    run_addr.font.name = "Arial"
    run_addr.font.size = Pt(10)
    run_addr.font.bold = True

    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_mob = p_info.add_run(f"Mob No: {data['vendor']['phone']}\t\t\t\t\tEmail Id: {data['vendor']['email'].upper()}\n")
    r_mob.font.size = Pt(9.5)
    r_mob.font.bold = True
    
    r_gst = p_info.add_run(f"TIN VAT No: 1056553025\t\t\t\t\tGSTIN: {data['vendor']['gstin']}")
    r_gst.font.size = Pt(9.5)
    r_gst.font.bold = True

    # Divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(12)
    p_div_border = p_div.paragraph_format
    # Add bottom border via OXML
    pPr = p_div._p.get_or_add_pPr()
    pbdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="12" w:space="1" w:color="000000"/></w:pBdr>')
    pPr.append(pbdr)

    # 2. Document Title
    p_doc_title = doc.add_paragraph()
    p_doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_doc_title = p_doc_title.add_run("WARRANTY / GUARANTEE CERTIFICATE")
    run_doc_title.font.name = "Arial"
    run_doc_title.font.size = Pt(13)
    run_doc_title.font.bold = True
    p_doc_title.paragraph_format.space_after = Pt(14)

    # 3. File No & Date
    gc_file_no = data.get("gc_file_no", "HT/GC-WC/26-27")
    gc_date = data.get("gc_date", data.get("challan_date", "24/07/2026"))
    
    p_file = doc.add_paragraph()
    p_file.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_file = p_file.add_run(f"File No.:- {gc_file_no}\t\t\t\t\tDate: {gc_date}")
    r_file.font.name = "Arial"
    r_file.font.size = Pt(10.5)
    r_file.font.bold = True
    p_file.paragraph_format.space_after = Pt(12)

    # 4. Dynamic "To" Block
    consignee_name = data.get("consignee", "SSE/DPS").strip()
    if not consignee_name.endswith('.'):
        consignee_name += "."

    p_to = doc.add_paragraph()
    p_to.paragraph_format.space_after = Pt(14)
    p_to.paragraph_format.line_spacing = 1.15
    
    r_to = p_to.add_run("To,\nThe,\n")
    r_to.font.name = "Arial"
    r_to.font.size = Pt(10.5)
    r_to.font.bold = True

    r_cons = p_to.add_run(f"{consignee_name}\n")
    r_cons.font.name = "Arial"
    r_cons.font.size = Pt(10.5)
    r_cons.font.bold = True

    r_rail = p_to.add_run("Eastern Railway, Jamalpur")
    r_rail.font.name = "Arial"
    r_rail.font.size = Pt(10.5)
    r_rail.font.bold = True

    # 5. Boxed Warranty Content
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    cell = table.cell(0, 0)
    cell.width = Inches(6.8)

    # Set thin black border for the cell
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        r'<w:left w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        r'<w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        r'<w:right w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        r'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

    # Cell Content
    cell_p = cell.paragraphs[0]
    cell_p.paragraph_format.space_after = Pt(8)
    cell_p.paragraph_format.line_spacing = 1.15

    r_cert = cell_p.add_run(
        "This is to certify that the materials supplied under the following purchase order are guaranteed as per IRS Terms & Conditions against any manufacturing defect or poor workmanship.\n\n"
    )
    r_cert.font.name = "Arial"
    r_cert.font.size = Pt(10)

    # PO & Challan Line
    po_no = data.get("po_number", "55265692101304")
    po_dt = data.get("po_date", "13/07/2026")
    ch_no = data.get("challan_no", "44")
    ch_dt = data.get("challan_date", "24/07/2026")

    p_po = cell.add_paragraph()
    p_po.paragraph_format.space_after = Pt(10)
    r_po = p_po.add_run(f"PO No.: {po_no}\t\tDated: {po_dt}\nDelivery Challan No.:{ch_no}\t\tDated: {ch_dt}")
    r_po.font.name = "Arial"
    r_po.font.size = Pt(10)
    r_po.font.bold = True

    # Material Description
    p_mat_head = cell.add_paragraph()
    p_mat_head.paragraph_format.space_after = Pt(6)
    r_mh = p_mat_head.add_run("Material Description:")
    r_mh.font.name = "Arial"
    r_mh.font.size = Pt(10)
    r_mh.font.bold = True

    for item in data.get("items", []):
        p_item = cell.add_paragraph()
        p_item.paragraph_format.space_after = Pt(6)
        
        desc_line = f"{item['sr_no']}. {item['description']}\n"
        r_desc = p_item.add_run(desc_line)
        r_desc.font.name = "Arial"
        r_desc.font.size = Pt(9.5)

        qty_str = item.get("quantity_display", f"{item['quantity']} Nos")
        r_qty = p_item.add_run(f"Quantity: {qty_str}")
        r_qty.font.name = "Arial"
        r_qty.font.size = Pt(9.5)
        r_qty.font.bold = True

    p_valid = cell.add_paragraph()
    p_valid.paragraph_format.space_before = Pt(12)
    p_valid.paragraph_format.space_after = Pt(8)
    r_valid = p_valid.add_run("The warranty/guarantee shall remain valid as per the IRS terms and conditions.")
    r_valid.font.name = "Arial"
    r_valid.font.size = Pt(9.5)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
